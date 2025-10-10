
# URL del artículo
# url = "https://www.dailymail.co.uk/tvshowbiz/article-13835203/my-unpleasant-encounter-with-kristina-rihanoff-and-ben-cohen-during-strictly-scandal-exposed-their-arrogance-and-i-cant-feel-sorry-for-them-now-theyve-fallen-writes-katie-hind.html?ns_mchannel=rss&ns_campaign=1490&ito=1490"

import requests
from bs4 import BeautifulSoup
from docx import Document

def scrape_and_save_to_word(url, filename="contenido_web.docx"):
    """
    Realiza web scraping de una URL y guarda el texto en un documento de Word.

    Args:
        url (str): La URL de la página web de la que se va a extraer el contenido.
        filename (str, optional): El nombre del archivo de Word que se creará. 
                                  Por defecto es "contenido_web.docx".
    """
    try:
        # Realizar la solicitud HTTP a la URL
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # Lanza un error para respuestas HTTP no exitosas

        # Analizar el contenido HTML de la página
        soup = BeautifulSoup(response.content, 'html.parser')

        # Crear un nuevo documento de Word
        document = Document()

        # Añadir el título de la página como un encabezado en el documento
        page_title = soup.title.string if soup.title else 'Sin Título'
        document.add_heading(page_title, level=1)

        # Extraer todo el texto de la página
        # El método get_text() extrae todo el texto de las etiquetas HTML
        # El parámetro 'strip=True' elimina los espacios en blanco innecesarios
        # El parámetro 'separator' añade un salto de línea entre los bloques de texto
        text = soup.get_text(separator='\n', strip=True)

        # Añadir el texto extraído al documento de Word
        document.add_paragraph(text)

        # Guardar el documento
        document.save(filename)
        print(f"¡El contenido de la web se ha guardado en '{filename}' exitosamente! ✅")

    except requests.exceptions.RequestException as e:
        print(f"Error al intentar acceder a la URL: {e}")
    except Exception as e:
        print(f"Ha ocurrido un error inesperado: {e}")


if __name__ == "__main__":
    # Pide al usuario que introduzca la URL
    # target_url = input("Introduce la URL de la página web que quieres descargar: ")
    target_url = "https://www.telegraph.co.uk/travel/terms-and-conditions/railbookers-terms-and-conditions/"
    # Llama a la función con la URL proporcionada
    scrape_and_save_to_word(target_url)
